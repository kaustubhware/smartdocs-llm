import streamlit as st
import os
from groq import Groq
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import pandas as pd
import plotly.express as px
from datetime import datetime
import json
import re
from collections import Counter
import difflib

# Load environment
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Page config
st.set_page_config(
    page_title="📊 SmartDocs",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;600;700&family=Open+Sans:wght@300;400;600;700&display=swap');
    
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        font-family: 'Roboto', sans-serif;
        color: #2c3e50;
    }
    
    .main-header {
        font-size: 2.8rem;
        font-weight: 600;
        text-align: center;
        color: #2c3e50;
        margin: 2rem 0 1rem 0;
        text-shadow: 0 2px 4px rgba(0,0,0,0.1);
        font-family: 'Open Sans', sans-serif;
    }
    
    .subtitle {
        text-align: center;
        color: #5a6c7d;
        font-size: 1.1rem;
        margin-bottom: 2rem;
        font-weight: 400;
        font-family: 'Open Sans', sans-serif;
    }
    
    .metric-card {
        background: #ffffff;
        padding: 1.8rem;
        border-radius: 8px;
        text-align: center;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border: 1px solid #e1e8ed;
        margin: 1rem 0;
        transition: all 0.2s ease;
    }
    
    .metric-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.12);
        transform: translateY(-2px);
    }
    
    .metric-number {
        font-size: 2.2rem;
        font-weight: 600;
        color: #2c3e50;
        margin-bottom: 0.5rem;
        font-family: 'Roboto', sans-serif;
    }
    
    .metric-label {
        color: #5a6c7d;
        font-weight: 500;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        font-size: 0.85rem;
    }
    
    .feature-card {
        background: #ffffff;
        padding: 2rem;
        border-radius: 8px;
        border: 1px solid #e1e8ed;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        margin: 1rem 0;
        transition: all 0.2s ease;
        text-align: center;
    }
    
    .feature-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.12);
        transform: translateY(-2px);
    }
    
    .feature-icon {
        font-size: 2.5rem;
        margin-bottom: 1rem;
        color: #3498db;
    }
    
    .feature-title {
        font-size: 1.3rem;
        font-weight: 600;
        color: #2c3e50;
        margin-bottom: 0.8rem;
        font-family: 'Open Sans', sans-serif;
    }
    
    .feature-desc {
        color: #5a6c7d;
        line-height: 1.5;
        font-size: 0.95rem;
    }
    
    .header-section {
        background: #ffffff;
        padding: 2rem;
        border-radius: 8px;
        margin-bottom: 2rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border: 1px solid #e1e8ed;
        text-align: center;
    }
    
    .footer {
        background: #2c3e50;
        color: #ffffff;
        padding: 1rem;
        text-align: center;
        margin-top: 3rem;
        border-radius: 8px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'processed' not in st.session_state:
    st.session_state.processed = False
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'language' not in st.session_state:
    st.session_state.language = 'English'

def extract_text(file):
    """Extract text from file with metadata"""
    try:
        if file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text, len(reader.pages)
        elif file.name.endswith('.docx'):
            doc = Document(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text, len(doc.paragraphs)
        elif file.name.endswith('.txt'):
            text = file.getvalue().decode('utf-8')
            return text, len(text.split('\n'))
        elif file.name.endswith(('.csv', '.xlsx')):
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
            return df.to_string(), len(df)
    except Exception as e:
        st.error(f"Error: {e}")
        return "", 0

def detect_language(text):
    """Simple language detection"""
    english_words = ['the', 'and', 'is', 'in', 'to', 'of', 'a', 'that', 'it', 'with']
    spanish_words = ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'es', 'se', 'no']
    french_words = ['le', 'de', 'et', 'à', 'un', 'il', 'être', 'et', 'en', 'avoir']
    marathi_words = ['आणि', 'हे', 'त्या', 'म्हणून', 'आहे', 'की', 'वर', 'ते', 'या', 'असे']
    
    text_lower = text.lower()
    english_count = sum(1 for word in english_words if word in text_lower)
    spanish_count = sum(1 for word in spanish_words if word in text_lower)
    french_count = sum(1 for word in french_words if word in text_lower)
    marathi_count = sum(1 for word in marathi_words if word in text)
    
    counts = {
        'English': english_count,
        'Spanish': spanish_count, 
        'French': french_count,
        'Marathi': marathi_count
    }
    
    return max(counts, key=counts.get)

def analyze_sentiment(text):
    """Simple sentiment analysis"""
    positive_words = ['good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'positive', 'success']
    negative_words = ['bad', 'terrible', 'awful', 'horrible', 'negative', 'failure', 'problem', 'issue']
    
    text_lower = text.lower()
    positive_count = sum(1 for word in positive_words if word in text_lower)
    negative_count = sum(1 for word in negative_words if word in text_lower)
    
    if positive_count > negative_count:
        return "Positive", positive_count / (positive_count + negative_count + 1)
    elif negative_count > positive_count:
        return "Negative", negative_count / (positive_count + negative_count + 1)
    else:
        return "Neutral", 0.5

def extract_key_topics(text, num_topics=5):
    """Extract key topics from text"""
    words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
    stop_words = {'this', 'that', 'with', 'have', 'will', 'from', 'they', 'been', 'were', 'said', 'each', 'which', 'their', 'time', 'would', 'there', 'could', 'other'}
    words = [word for word in words if word not in stop_words]
    
    word_freq = Counter(words)
    return word_freq.most_common(num_topics)

def estimate_reading_time(text):
    """Estimate reading time (average 200 words per minute)"""
    word_count = len(text.split())
    reading_time = word_count / 200
    return max(1, round(reading_time))

def chat_with_groq(question, context, language="English"):
    """Enhanced chat with language support"""
    if not GROQ_API_KEY:
        return "Please add GROQ_API_KEY to .env file"
    
    client = Groq(api_key=GROQ_API_KEY)
    
    language_prompts = {
        "English": "Answer in English based on the provided context.",
        "Spanish": "Responde en español basándote en el contexto proporcionado.",
        "French": "Répondez en français basé sur le contexte fourni.",
        "German": "Antworten Sie auf Deutsch basierend auf dem bereitgestellten Kontext.",
        "Italian": "Rispondi in italiano basandoti sul contesto fornito.",
        "Marathi": "प्रदान केलेल्या संदर्भावर आधारित मराठीत उत्तर द्या."
    }
    
    system_prompt = language_prompts.get(language, language_prompts["English"])
    
    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Context: {context}\n\nQuestion: {question}"}
            ],
            temperature=0.1
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {e}"

# Main UI
st.markdown("""
<div class="header-section">
    <h1 class="main-header">📊 SmartDocs</h1>
    <p class="subtitle">Advanced AI-Powered Document Analysis & Multi-Language Processing Solution</p>
</div>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.header("🎛️ Control Panel")
    
    # Language selection
    st.session_state.language = st.selectbox(
        "🌍 Response Language",
        ["English", "Spanish", "French", "German", "Italian", "Marathi"]
    )
    
    st.header("📁 Document Management")
    
    uploaded_files = st.file_uploader(
        "Upload Documents",
        type=['pdf', 'docx', 'txt', 'csv', 'xlsx'],
        accept_multiple_files=True,
        help="Supports: PDF, DOCX, TXT, CSV, XLSX"
    )
    
    if uploaded_files and st.button("🚀 Process Documents", type="primary"):
        st.session_state.documents = []
        
        progress_bar = st.progress(0)
        for i, file in enumerate(uploaded_files):
            text, metadata = extract_text(file)
            if text:
                language = detect_language(text)
                sentiment, sentiment_score = analyze_sentiment(text)
                topics = extract_key_topics(text)
                reading_time = estimate_reading_time(text)
                
                st.session_state.documents.append({
                    'name': file.name,
                    'content': text,
                    'language': language,
                    'sentiment': sentiment,
                    'sentiment_score': sentiment_score,
                    'topics': topics,
                    'reading_time': reading_time,
                    'word_count': len(text.split()),
                    'metadata': metadata
                })
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        if st.session_state.documents:
            st.session_state.processed = True
            st.success(f"✅ Processed {len(st.session_state.documents)} documents")
            st.rerun()
    
    # Quick stats
    if st.session_state.documents:
        st.header("📊 Quick Stats")
        total_words = sum(doc['word_count'] for doc in st.session_state.documents)
        total_reading_time = sum(doc['reading_time'] for doc in st.session_state.documents)
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 Files", len(st.session_state.documents))
            st.metric("📖 Reading Time", f"{total_reading_time} min")
        with col2:
            st.metric("📝 Words", f"{total_words:,}")
            languages = set(doc['language'] for doc in st.session_state.documents)
            st.metric("🌍 Languages", len(languages))

# Main content
if not st.session_state.processed:
    st.info("👆 Upload and process documents to unlock all features!")
    
    st.header("🌟 Advanced Features")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🤖</div>
            <div class="feature-title">Multi-Language AI</div>
            <div class="feature-desc">Chat in 6+ languages with automatic detection and intelligent responses</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">📊</div>
            <div class="feature-title">Smart Analytics</div>
            <div class="feature-desc">Advanced sentiment analysis, topic extraction, and interactive visualizations</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🔍</div>
            <div class="feature-title">Document Comparison</div>
            <div class="feature-desc">AI-powered side-by-side analysis with similarity scoring and topic mapping</div>
        </div>
        """, unsafe_allow_html=True)

else:
    # Main tabs
    tab1, tab2, tab3, tab4 = st.tabs(["💬 Smart Chat", "📊 Analytics", "🔍 Document Comparison", "📥 Export"])
    
    with tab1:
        st.header("💬 Smart Chat")
        
        # Chat history
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                st.chat_message("user").write(msg['content'])
            else:
                st.chat_message("assistant").write(msg['content'])
        
        # Chat input
        if 'last_prompt' not in st.session_state:
            st.session_state.last_prompt = ""
            
        prompt = st.text_input(f"Ask in {st.session_state.language}...", key="chat_input")
        
        if prompt and prompt != st.session_state.last_prompt:
            st.session_state.last_prompt = prompt
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            st.chat_message("user").write(prompt)
            
            # Combine all documents
            context = "\n\n".join([f"Document: {doc['name']}\n{doc['content']}" for doc in st.session_state.documents])
            
            with st.spinner("🤖 Generating intelligent response..."):
                response = chat_with_groq(prompt, context[:8000], st.session_state.language)
            
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            st.chat_message("assistant").write(response)
            st.rerun()
    
    with tab2:
        st.header("📊 Advanced Analytics")
        
        if st.session_state.documents:
            # Overview metrics
            col1, col2, col3, col4 = st.columns(4)
            
            total_words = sum(doc['word_count'] for doc in st.session_state.documents)
            total_reading_time = sum(doc['reading_time'] for doc in st.session_state.documents)
            languages = list(set(doc['language'] for doc in st.session_state.documents))
            avg_sentiment = sum(doc['sentiment_score'] for doc in st.session_state.documents) / len(st.session_state.documents)
            
            with col1:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-number">{len(st.session_state.documents)}</div>
                    <div class="metric-label">Documents</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-number">{total_words:,}</div>
                    <div class="metric-label">Total Words</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-number">{total_reading_time}</div>
                    <div class="metric-label">Minutes to Read</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col4:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-number">{avg_sentiment:.2f}</div>
                    <div class="metric-label">Avg Sentiment</div>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("---")
            
            # Charts
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📈 Document Statistics")
                
                # Word count chart
                doc_data = pd.DataFrame([
                    {'Document': doc['name'], 'Words': doc['word_count'], 'Reading Time': doc['reading_time']}
                    for doc in st.session_state.documents
                ])
                
                fig = px.bar(doc_data, x='Document', y='Words', title='Word Count by Document')
                fig.update_layout(xaxis_tickangle=-45)
                st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                st.subheader("🌍 Language Distribution")
                
                # Language distribution
                lang_counts = Counter(doc['language'] for doc in st.session_state.documents)
                lang_df = pd.DataFrame(list(lang_counts.items()), columns=['Language', 'Count'])
                
                fig = px.pie(lang_df, values='Count', names='Language', title='Documents by Language')
                st.plotly_chart(fig, use_container_width=True)
            
            # Sentiment analysis
            st.subheader("😊 Sentiment Analysis")
            
            sentiment_data = pd.DataFrame([
                {'Document': doc['name'], 'Sentiment': doc['sentiment'], 'Score': doc['sentiment_score']}
                for doc in st.session_state.documents
            ])
            
            fig = px.bar(sentiment_data, x='Document', y='Score', color='Sentiment',
                        title='Sentiment Analysis by Document')
            fig.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig, use_container_width=True)
            
            # Top topics across all documents
            st.subheader("🏷️ Top Topics Across All Documents")
            
            all_topics = []
            for doc in st.session_state.documents:
                all_topics.extend(doc['topics'])
            
            topic_counter = Counter()
            for topic, count in all_topics:
                topic_counter[topic] += count
            
            top_topics = topic_counter.most_common(10)
            
            if top_topics:
                topics_df = pd.DataFrame(top_topics, columns=['Topic', 'Frequency'])
                fig = px.bar(topics_df, x='Frequency', y='Topic', orientation='h',
                           title='Most Frequent Topics')
                st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("📊 Upload documents to see analytics")
    
    with tab3:
        st.header("🔍 Document Comparison")
        
        if len(st.session_state.documents) >= 2:
            col1, col2 = st.columns(2)
            
            with col1:
                doc1_name = st.selectbox("Select First Document", 
                                       [doc['name'] for doc in st.session_state.documents],
                                       key="doc1")
            
            with col2:
                doc2_name = st.selectbox("Select Second Document", 
                                       [doc['name'] for doc in st.session_state.documents],
                                       key="doc2")
            
            if doc1_name != doc2_name:
                doc1 = next(doc for doc in st.session_state.documents if doc['name'] == doc1_name)
                doc2 = next(doc for doc in st.session_state.documents if doc['name'] == doc2_name)
                
                # Similarity calculation
                similarity = difflib.SequenceMatcher(None, doc1['content'], doc2['content']).ratio()
                
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-number">{similarity:.2%}</div>
                    <div class="metric-label">Similarity Score</div>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader(f"📄 {doc1_name}")
                    st.write(f"**Language:** {doc1['language']}")
                    st.write(f"**Sentiment:** {doc1['sentiment']} ({doc1['sentiment_score']:.2f})")
                    st.write(f"**Word Count:** {doc1['word_count']:,}")
                    st.write(f"**Reading Time:** {doc1['reading_time']} min")
                    
                    st.write("**Top Topics:**")
                    for topic, count in doc1['topics'][:5]:
                        st.write(f"• {topic} ({count})")
                
                with col2:
                    st.subheader(f"📄 {doc2_name}")
                    st.write(f"**Language:** {doc2['language']}")
                    st.write(f"**Sentiment:** {doc2['sentiment']} ({doc2['sentiment_score']:.2f})")
                    st.write(f"**Word Count:** {doc2['word_count']:,}")
                    st.write(f"**Reading Time:** {doc2['reading_time']} min")
                    
                    st.write("**Top Topics:**")
                    for topic, count in doc2['topics'][:5]:
                        st.write(f"• {topic} ({count})")
                
                # Common and unique topics
                st.subheader("🏷️ Topic Analysis")
                
                doc1_topics = set(topic for topic, _ in doc1['topics'])
                doc2_topics = set(topic for topic, _ in doc2['topics'])
                
                common_topics = doc1_topics.intersection(doc2_topics)
                unique_doc1 = doc1_topics - doc2_topics
                unique_doc2 = doc2_topics - doc1_topics
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.write("**Common Topics:**")
                    for topic in list(common_topics)[:5]:
                        st.write(f"• {topic}")
                
                with col2:
                    st.write(f"**Unique to {doc1_name}:**")
                    for topic in list(unique_doc1)[:5]:
                        st.write(f"• {topic}")
                
                with col3:
                    st.write(f"**Unique to {doc2_name}:**")
                    for topic in list(unique_doc2)[:5]:
                        st.write(f"• {topic}")
        
        elif len(st.session_state.documents) == 1:
            st.info("📄 Upload at least 2 documents to compare")
        else:
            st.info("📄 Upload documents to enable comparison")
    
    with tab4:
        st.header("📥 Export & Reports")
        
        if st.session_state.documents:
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("💾 Export Options")
                
                # Export chat history
                if st.session_state.chat_history:
                    chat_json = json.dumps(st.session_state.chat_history, indent=2)
                    st.download_button(
                        "📝 Download Chat History (JSON)",
                        chat_json,
                        f"chat_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        "application/json"
                    )
                
                # Export analytics data
                analytics_data = {
                    'summary': {
                        'total_documents': len(st.session_state.documents),
                        'total_words': sum(doc['word_count'] for doc in st.session_state.documents),
                        'total_reading_time': sum(doc['reading_time'] for doc in st.session_state.documents),
                        'languages': list(set(doc['language'] for doc in st.session_state.documents))
                    },
                    'documents': st.session_state.documents
                }
                
                analytics_json = json.dumps(analytics_data, indent=2, default=str)
                st.download_button(
                    "📊 Download Analytics Data (JSON)",
                    analytics_json,
                    f"analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    "application/json"
                )
            
            with col2:
                st.subheader("📋 Summary Report")
                
                total_words = sum(doc['word_count'] for doc in st.session_state.documents)
                total_reading_time = sum(doc['reading_time'] for doc in st.session_state.documents)
                languages = list(set(doc['language'] for doc in st.session_state.documents))
                
                report = f"""
# Document Analysis Report

## Summary
- **Total Documents:** {len(st.session_state.documents)}
- **Total Words:** {total_words:,}
- **Estimated Reading Time:** {total_reading_time} minutes
- **Languages Detected:** {', '.join(languages)}

## Document Details

"""
                
                for doc in st.session_state.documents:
                    report += f"""
### {doc['name']}
- **Language:** {doc['language']}
- **Sentiment:** {doc['sentiment']} (Score: {doc['sentiment_score']:.2f})
- **Word Count:** {doc['word_count']:,}
- **Reading Time:** {doc['reading_time']} minutes
- **Top Topics:** {', '.join([topic for topic, _ in doc['topics'][:3]])}

"""
                
                st.markdown(report)
                
                st.download_button(
                    "📄 Download Report (Markdown)",
                    report,
                    f"document_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    "text/markdown"
                )
        else:
            st.info("📄 Upload documents to enable export features")

# Footer
st.markdown("""
<div class="footer">
    📊 SmartDocs - Powered by Advanced AI Technology
</div>
""", unsafe_allow_html=True)
